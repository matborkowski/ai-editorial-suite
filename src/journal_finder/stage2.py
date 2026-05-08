from __future__ import annotations

import argparse
import csv
import math
import re
import sys
import time
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity

    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False


DEFAULT_OUTPUT_SUFFIX = "JournalFinder_results.docx"
DEFAULT_TIMEOUT = 12
DEFAULT_MAX_PROFILE_CHARS = 9000
DEFAULT_WEB_DELAY_SECONDS = 0.5
DEFAULT_TOP_N = 3

STAGE1_REQUIRED_COLUMNS = {
    "tytul czasopisma": "title",
    "liczba punktow": "points",
    "dyscypliny naukowe": "disciplines",
    "issn": "issn",
}

PROFILE_LINK_PATTERNS = [
    "aims and scope",
    "aims & scope",
    "aims/scope",
    "scope",
    "about this journal",
    "about the journal",
    "journal overview",
    "overview",
    "description",
]

PROFILE_TEXT_PATTERNS = [
    "aims and scope",
    "aims & scope",
    "about this journal",
    "about the journal",
    "journal overview",
    "scope",
    "description",
    "publishes",
    "journal publishes",
]

EN_STOPWORDS = {
    "a", "an", "and", "are", "as", "at", "be", "been", "by", "for", "from", "has", "have", "in", "into",
    "is", "it", "its", "of", "on", "or", "that", "the", "their", "this", "to", "using", "with", "within",
    "without", "we", "were", "was", "will", "can", "could", "may", "might", "study", "paper", "article",
}

PL_STOPWORDS = {
    "a", "aby", "albo", "ale", "bez", "byc", "czy", "dla", "do", "i", "ich", "jak", "jest", "jako", "lub",
    "na", "nad", "nie", "oraz", "po", "pod", "przez", "sa", "sie", "tak", "to", "w", "we", "z", "za", "ze",
}

FILENAME_STOPWORDS = EN_STOPWORDS | {"the", "of", "and", "in", "on", "for", "with", "using"}


class JournalFinderStage2Error(Exception):
    pass


@dataclass
class JournalCandidate:
    title: str
    points: str = ""
    disciplines: str = ""
    issn: str = ""
    profile_url: str = ""
    profile_text: str = ""
    impact_factor: str = "brak danych"
    snip: str = "brak danych"
    citescore: str = "brak danych"
    source_status: str = "brak profilu"
    score: float = 1.0
    confidence: str = "niska"
    matched_keywords: list[str] = field(default_factory=list)
    justification: str = ""


@dataclass
class MetadataRecord:
    profile_url: str = ""
    profile_text: str = ""
    impact_factor: str = "brak danych"
    snip: str = "brak danych"
    citescore: str = "brak danych"


def normalize_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value).strip().casefold()
    # NFKD nie rozkłada polskiego „ł”, więc trzeba obsłużyć je jawnie.
    text = text.replace("ł", "l")
    text = unicodedata.normalize("NFKD", text)
    text = "".join(char for char in text if not unicodedata.combining(char))
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_header(value: object) -> str:
    text = normalize_text(value)
    text = re.sub(r"[^a-z0-9 ]+", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def clean_cell(value: object) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    if text.lower() in {"nan", "none", "null"}:
        return ""
    return text


def normalize_issn(value: object) -> str:
    text = clean_cell(value).upper()
    text = re.sub(r"[^0-9X]", "", text)
    return text


def split_issns(value: object) -> list[str]:
    raw = clean_cell(value)
    if not raw:
        return []
    parts = re.split(r"[;,/|\s]+", raw)
    result: list[str] = []
    for part in parts:
        norm = normalize_issn(part)
        if len(norm) == 8 and norm not in result:
            result.append(norm)
    # Awaryjnie: wyszukaj ISSN w długim tekście, np. 1234-567X.
    for match in re.findall(r"\b\d{4}[- ]?\d{3}[0-9Xx]\b", raw):
        norm = normalize_issn(match)
        if len(norm) == 8 and norm not in result:
            result.append(norm)
    return result


def tokenize(text: str) -> list[str]:
    normalized = normalize_text(text)
    tokens = re.findall(r"\b[a-z0-9][a-z0-9\-]{2,}\b", normalized)
    stopwords = EN_STOPWORDS | PL_STOPWORDS
    return [token for token in tokens if token not in stopwords and not token.isdigit()]


def unique_keep_order(values: Iterable[str]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for value in values:
        key = normalize_text(value)
        if key and key not in seen:
            seen.add(key)
            result.append(value)
    return result


def make_output_filename(article_title: str) -> str:
    tokens = re.findall(r"[A-Za-z0-9]+", unicodedata.normalize("NFKD", article_title))
    cleaned: list[str] = []
    for token in tokens:
        norm = normalize_text(token)
        if norm and norm not in FILENAME_STOPWORDS:
            safe = re.sub(r"[^A-Za-z0-9]+", "", token)
            if safe:
                cleaned.append(safe)
        if len(cleaned) == 3:
            break
    if not cleaned:
        cleaned = ["Article"]
    return "_".join(cleaned) + "_" + DEFAULT_OUTPUT_SUFFIX


def read_text_argument(value: str | None, file_path: str | None, field_name: str) -> str:
    if value and file_path:
        raise JournalFinderStage2Error(f"Podaj {field_name} albo jako tekst, albo jako plik — nie oba naraz.")
    if file_path:
        path = Path(file_path)
        if not path.exists():
            raise JournalFinderStage2Error(f"Nie znaleziono pliku z polem {field_name}: {path}")
        return path.read_text(encoding="utf-8").strip()
    return clean_cell(value)


def load_stage1_docx(stage1_docx_path: Path) -> list[JournalCandidate]:
    if not stage1_docx_path.exists():
        raise JournalFinderStage2Error(f"Nie znaleziono pliku Stage 1 DOCX: {stage1_docx_path}")

    document = Document(stage1_docx_path)
    candidates: list[JournalCandidate] = []

    for table in document.tables:
        if not table.rows:
            continue

        headers = [normalize_header(cell.text) for cell in table.rows[0].cells]
        mapped_indices: dict[str, int] = {}
        for idx, header in enumerate(headers):
            if header in STAGE1_REQUIRED_COLUMNS:
                mapped_indices[STAGE1_REQUIRED_COLUMNS[header]] = idx

        if not {"title", "points", "disciplines", "issn"}.issubset(mapped_indices):
            continue

        for row in table.rows[1:]:
            cells = row.cells
            title = clean_cell(cells[mapped_indices["title"]].text)
            if not title:
                continue
            candidate = JournalCandidate(
                title=title,
                points=clean_cell(cells[mapped_indices["points"]].text),
                disciplines=clean_cell(cells[mapped_indices["disciplines"]].text),
                issn=clean_cell(cells[mapped_indices["issn"]].text),
            )
            candidates.append(candidate)

    if not candidates:
        raise JournalFinderStage2Error(
            "Nie znaleziono tabeli wynikowej Stage 1. Sprawdź, czy DOCX pochodzi z Journal Finder Stage 1."
        )

    return candidates


def load_stage1_docx_files(stage1_docx_paths: list[Path]) -> list[JournalCandidate]:
    all_candidates: list[JournalCandidate] = []
    seen: set[tuple[str, str]] = set()

    for path in stage1_docx_paths:
        for candidate in load_stage1_docx(path):
            key = (normalize_text(candidate.title), normalize_text(candidate.issn))
            if key not in seen:
                seen.add(key)
                all_candidates.append(candidate)

    if not all_candidates:
        raise JournalFinderStage2Error("Nie znaleziono żadnych czasopism w plikach Stage 1.")
    return all_candidates


def load_metadata_csv(path: Path | None) -> tuple[dict[str, MetadataRecord], dict[str, MetadataRecord]]:
    """
    CSV opcjonalny. Obsługiwane kolumny:
    title / tytuł czasopisma, ISSN, profile_url, profile_text, IF, SNIP, CiteScore.
    Zwraca mapy po ISSN i po tytule.
    """
    by_issn: dict[str, MetadataRecord] = {}
    by_title: dict[str, MetadataRecord] = {}

    if path is None:
        return by_issn, by_title
    if not path.exists():
        raise JournalFinderStage2Error(f"Nie znaleziono pliku metadata CSV: {path}")

    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        if not reader.fieldnames:
            raise JournalFinderStage2Error("Plik metadata CSV nie ma nagłówków.")

        normalized_fields = {normalize_header(name): name for name in reader.fieldnames}

        def get(row: dict[str, str], *names: str) -> str:
            for name in names:
                source_name = normalized_fields.get(normalize_header(name))
                if source_name and clean_cell(row.get(source_name)):
                    return clean_cell(row.get(source_name))
            return ""

        for row in reader:
            title = get(row, "title", "tytuł czasopisma", "tytul czasopisma", "journal title")
            issn = get(row, "ISSN", "issn")
            record = MetadataRecord(
                profile_url=get(row, "profile_url", "url", "official_url", "aims_scope_url"),
                profile_text=get(row, "profile_text", "aims_scope", "aims and scope", "profil", "scope"),
                impact_factor=get(row, "IF", "impact_factor", "impact factor") or "brak danych",
                snip=get(row, "SNIP", "snip") or "brak danych",
                citescore=get(row, "CiteScore", "citescore", "cite score") or "brak danych",
            )

            for issn_norm in split_issns(issn):
                by_issn[issn_norm] = record
            if title:
                by_title[normalize_text(title)] = record

    return by_issn, by_title


def apply_metadata(candidate: JournalCandidate, by_issn: dict[str, MetadataRecord], by_title: dict[str, MetadataRecord]) -> None:
    record: MetadataRecord | None = None
    for issn in split_issns(candidate.issn):
        if issn in by_issn:
            record = by_issn[issn]
            break
    if record is None:
        record = by_title.get(normalize_text(candidate.title))

    if record is None:
        return

    candidate.profile_url = record.profile_url or candidate.profile_url
    candidate.profile_text = record.profile_text or candidate.profile_text
    candidate.impact_factor = record.impact_factor or "brak danych"
    candidate.snip = record.snip or "brak danych"
    candidate.citescore = record.citescore or "brak danych"
    if candidate.profile_text:
        candidate.source_status = "profil z CSV"


def build_session() -> requests.Session:
    session = requests.Session()
    session.headers.update(
        {
            "User-Agent": "JournalFinderStage2/1.0 (+local research tool; contact: user)",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        }
    )
    return session


def fetch_url(session: requests.Session, url: str, timeout: int = DEFAULT_TIMEOUT) -> tuple[str, str]:
    response = session.get(url, timeout=timeout, allow_redirects=True)
    response.raise_for_status()
    content_type = response.headers.get("content-type", "")
    if "text/html" not in content_type and "application/xhtml" not in content_type and not response.text.lstrip().startswith("<"):
        raise JournalFinderStage2Error(f"URL nie wygląda na stronę HTML: {url}")
    return response.url, response.text


def html_to_text(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg", "form"]):
        tag.decompose()
    text = soup.get_text(" ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_profile_text(page_text: str, max_chars: int = DEFAULT_MAX_PROFILE_CHARS) -> str:
    if not page_text:
        return ""
    normalized = normalize_text(page_text)
    positions: list[int] = []
    for pattern in PROFILE_TEXT_PATTERNS:
        idx = normalized.find(pattern)
        if idx >= 0:
            positions.append(idx)
    if positions:
        start = max(0, min(positions) - 800)
        end = min(len(page_text), start + max_chars)
        return page_text[start:end].strip()
    return page_text[:max_chars].strip()


def same_registered_domain(url_a: str, url_b: str) -> bool:
    host_a = urlparse(url_a).netloc.lower().removeprefix("www.")
    host_b = urlparse(url_b).netloc.lower().removeprefix("www.")
    if not host_a or not host_b:
        return False
    return host_a == host_b or host_a.endswith("." + host_b) or host_b.endswith("." + host_a)


def discover_profile_links(base_url: str, html: str, limit: int = 5) -> list[str]:
    soup = BeautifulSoup(html, "html.parser")
    candidates: list[tuple[int, str]] = []
    for a_tag in soup.find_all("a", href=True):
        label = normalize_text(a_tag.get_text(" "))
        href = clean_cell(a_tag.get("href"))
        if not href:
            continue
        absolute = urljoin(base_url, href)
        if not same_registered_domain(base_url, absolute):
            continue
        score = 0
        for pattern in PROFILE_LINK_PATTERNS:
            if pattern in label or pattern.replace(" ", "-") in normalize_text(href):
                score += 1
        if score > 0:
            candidates.append((score, absolute))

    candidates.sort(key=lambda item: item[0], reverse=True)
    return unique_keep_order([url for _, url in candidates])[:limit]


def profile_quality_score(text: str) -> int:
    normalized = normalize_text(text)
    return sum(1 for pattern in PROFILE_TEXT_PATTERNS if pattern in normalized)


def fetch_official_profile(session: requests.Session, url: str, web_delay: float = DEFAULT_WEB_DELAY_SECONDS) -> tuple[str, str, str]:
    """Zwraca: final_url, profile_text, status."""
    final_url, html = fetch_url(session, url)
    home_text = html_to_text(html)
    best_url = final_url
    best_text = extract_profile_text(home_text)
    best_quality = profile_quality_score(best_text)

    for profile_link in discover_profile_links(final_url, html):
        time.sleep(web_delay)
        try:
            candidate_final_url, candidate_html = fetch_url(session, profile_link)
            candidate_text = extract_profile_text(html_to_text(candidate_html))
            candidate_quality = profile_quality_score(candidate_text)
            if candidate_quality > best_quality or (candidate_quality == best_quality and len(candidate_text) > len(best_text)):
                best_url = candidate_final_url
                best_text = candidate_text
                best_quality = candidate_quality
        except Exception:
            continue

    status = "oficjalny profil pobrany" if best_text else "brak tekstu profilu"
    return best_url, best_text, status


def discover_homepage_with_openalex(session: requests.Session, candidate: JournalCandidate) -> str:
    """
    OpenAlex służy tylko do znalezienia homepage_url po ISSN.
    Profil/scope nadal musi zostać pobrany z docelowej strony czasopisma/wydawcy.
    """
    for issn in split_issns(candidate.issn):
        api_url = f"https://api.openalex.org/sources/issn:{issn}"
        try:
            response = session.get(api_url, timeout=DEFAULT_TIMEOUT)
            if response.status_code == 404:
                continue
            response.raise_for_status()
            data = response.json()
            homepage_url = clean_cell(data.get("homepage_url"))
            if homepage_url:
                return homepage_url
        except Exception:
            continue
    return ""


def enrich_candidate_with_web(
    session: requests.Session,
    candidate: JournalCandidate,
    use_openalex_discovery: bool,
    web_delay: float,
) -> None:
    if candidate.profile_text:
        return

    if not candidate.profile_url and use_openalex_discovery:
        candidate.profile_url = discover_homepage_with_openalex(session, candidate)
        if candidate.profile_url:
            candidate.source_status = "URL znaleziony przez OpenAlex"

    if not candidate.profile_url:
        candidate.source_status = "brak oficjalnego URL"
        return

    try:
        final_url, profile_text, status = fetch_official_profile(session, candidate.profile_url, web_delay=web_delay)
        candidate.profile_url = final_url
        candidate.profile_text = profile_text
        candidate.source_status = status
    except Exception as exc:
        candidate.source_status = f"błąd pobierania profilu: {type(exc).__name__}"


def keyword_matches(article_keywords: list[str], journal_text: str) -> list[str]:
    normalized_journal = normalize_text(journal_text)
    matches: list[str] = []
    for keyword in article_keywords:
        keyword_norm = normalize_text(keyword)
        if keyword_norm and keyword_norm in normalized_journal:
            matches.append(keyword)
    return matches


def fallback_similarity(article_text: str, journal_text: str) -> float:
    article_tokens = set(tokenize(article_text))
    journal_tokens = set(tokenize(journal_text))
    if not article_tokens or not journal_tokens:
        return 0.0
    common = article_tokens & journal_tokens
    coverage = len(common) / max(1, len(article_tokens))
    jaccard = len(common) / max(1, len(article_tokens | journal_tokens))
    return 0.75 * coverage + 0.25 * jaccard


def calculate_scores(candidates: list[JournalCandidate], article_text: str, article_keywords: list[str]) -> None:
    journal_documents: list[str] = []
    for candidate in candidates:
        # Profil jest najważniejszy. Tytuł i dyscypliny pomagają, ale nie powinny dominować.
        document = "\n".join(
            [
                candidate.profile_text,
                candidate.title,
                candidate.disciplines,
            ]
        ).strip()
        journal_documents.append(document)

    similarities: list[float]
    if SKLEARN_AVAILABLE and any(journal_documents):
        vectorizer = TfidfVectorizer(
            lowercase=True,
            strip_accents="unicode",
            stop_words="english",
            ngram_range=(1, 2),
            min_df=1,
        )
        matrix = vectorizer.fit_transform([article_text] + journal_documents)
        similarities = cosine_similarity(matrix[0:1], matrix[1:]).flatten().tolist()
    else:
        similarities = [fallback_similarity(article_text, doc) for doc in journal_documents]

    for candidate, similarity, journal_doc in zip(candidates, similarities, journal_documents):
        candidate.matched_keywords = keyword_matches(article_keywords, journal_doc)
        keyword_coverage = len(candidate.matched_keywords) / max(1, len(article_keywords))

        has_profile = bool(candidate.profile_text and len(candidate.profile_text) >= 300)
        profile_cap = 10.0 if has_profile else 5.0

        # Cosine TF-IDF dla krótkiego artykułu vs profil często mieści się w zakresie 0.00-0.25.
        similarity_component = min(similarity / 0.25, 1.0)
        raw_score = 1.0 + 7.0 * similarity_component + 2.0 * keyword_coverage
        if not has_profile:
            raw_score = min(raw_score, profile_cap)

        candidate.score = round(max(1.0, min(10.0, raw_score)), 1)

        if has_profile and candidate.score >= 7:
            candidate.confidence = "wysoka"
        elif has_profile and candidate.score >= 4:
            candidate.confidence = "średnia"
        else:
            candidate.confidence = "niska"

        candidate.justification = build_justification(candidate, similarity, has_profile)


def build_justification(candidate: JournalCandidate, similarity: float, has_profile: bool) -> str:
    parts: list[str] = []
    if has_profile:
        parts.append("Ocena oparta na profilu/aims & scope czasopisma oraz zgodności terminologicznej z artykułem.")
    else:
        parts.append("Ocena wstępna: brak pełnego oficjalnego profilu, więc wynik opiera się głównie na tytule czasopisma i dyscyplinach.")

    if candidate.matched_keywords:
        parts.append("Dopasowane słowa kluczowe: " + ", ".join(candidate.matched_keywords[:8]) + ".")
    else:
        parts.append("Nie znaleziono bezpośrednich trafień dla podanych słów kluczowych.")

    if candidate.score <= 3:
        parts.append("Profil artykułu wydaje się słabo zgodny z zakresem czasopisma albo brakuje danych do rzetelnej oceny.")
    elif candidate.score >= 7:
        parts.append("Zgodność tematyczna wygląda obiecująco, ale wynik nadal wymaga ręcznej weryfikacji przez redaktora/autora.")
    else:
        parts.append("Dopasowanie jest umiarkowane; czasopismo może być opcją rezerwową po sprawdzeniu szczegółowych wymagań.")

    parts.append(f"Podobieństwo tekstowe: {similarity:.3f}.")
    return " ".join(parts)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(8)


def add_key_value_paragraph(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(label).bold = True
    paragraph.add_run(value)


def add_table(document: Document, rows: list[list[str]], widths: list[Cm] | None = None) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col_idx, header in enumerate(rows[0]):
        cell = table.rows[0].cells[col_idx]
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row_values in rows[1:]:
        cells = table.add_row().cells
        for col_idx, value in enumerate(row_values):
            cells[col_idx].text = clean_cell(value)
            cells[col_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths[: len(row.cells)]):
                row.cells[idx].width = width


def export_results_docx(
    candidates: list[JournalCandidate],
    article_title: str,
    abstract: str,
    keywords: list[str],
    output_path: Path,
    top_n: int = DEFAULT_TOP_N,
) -> None:
    sorted_candidates = sorted(candidates, key=lambda item: item.score, reverse=True)
    top_candidates = sorted_candidates[:top_n]

    document = Document()
    configure_document(document)

    document.add_heading("Journal Finder – Stage 2 results", level=1)
    add_key_value_paragraph(document, "Tytuł artykułu: ", article_title)
    add_key_value_paragraph(document, "Słowa kluczowe: ", "; ".join(keywords))
    add_key_value_paragraph(document, "Liczba czasopism z Stage 1: ", str(len(candidates)))

    document.add_heading("Trzy najlepiej dopasowane czasopisma", level=2)
    top_rows = [["Miejsce", "Tytuł czasopisma", "Punkty", "ISSN", "Ocena", "IF", "SNIP", "CiteScore", "Uzasadnienie"]]
    for idx, candidate in enumerate(top_candidates, start=1):
        top_rows.append(
            [
                str(idx),
                candidate.title,
                candidate.points,
                candidate.issn,
                f"{candidate.score}/10",
                candidate.impact_factor,
                candidate.snip,
                candidate.citescore,
                candidate.justification,
            ]
        )
    add_table(
        document,
        top_rows,
        widths=[Cm(1.4), Cm(5.2), Cm(1.5), Cm(2.8), Cm(1.6), Cm(1.5), Cm(1.5), Cm(1.8), Cm(14.0)],
    )

    document.add_heading("Pełna tabela dopasowania", level=2)
    rows = [[
        "Lp.",
        "Tytuł czasopisma",
        "Punkty",
        "Dyscypliny naukowe",
        "ISSN",
        "IF",
        "SNIP",
        "CiteScore",
        "Ocena",
        "Pewność",
        "Źródło profilu",
        "URL",
        "Uzasadnienie",
    ]]
    for idx, candidate in enumerate(sorted_candidates, start=1):
        rows.append(
            [
                str(idx),
                candidate.title,
                candidate.points,
                candidate.disciplines,
                candidate.issn,
                candidate.impact_factor,
                candidate.snip,
                candidate.citescore,
                f"{candidate.score}/10",
                candidate.confidence,
                candidate.source_status,
                candidate.profile_url or "brak danych",
                candidate.justification,
            ]
        )
    add_table(
        document,
        rows,
        widths=[
            Cm(1.0), Cm(4.2), Cm(1.4), Cm(5.5), Cm(2.5), Cm(1.3), Cm(1.3), Cm(1.6), Cm(1.5), Cm(1.6),
            Cm(2.8), Cm(4.5), Cm(8.0),
        ],
    )

    document.add_page_break()
    document.add_heading("Dane wejściowe artykułu", level=2)
    document.add_heading("Abstrakt", level=3)
    document.add_paragraph(abstract)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def parse_keywords(raw_keywords: str) -> list[str]:
    keywords = [part.strip() for part in re.split(r"[;,\n]+", raw_keywords or "") if part.strip()]
    return unique_keep_order(keywords)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Journal Finder Stage 2 — dobór czasopisma do tytułu, abstraktu i słów kluczowych.")
    parser.add_argument(
        "--stage1-docx",
        required=True,
        nargs="+",
        help="Jeden lub kilka plików DOCX wygenerowanych w Stage 1, np. part1 part2 part3.",
    )
    parser.add_argument("--article-title", required=True, help="Tytuł artykułu.")
    parser.add_argument("--abstract", default=None, help="Abstrakt jako tekst.")
    parser.add_argument("--abstract-file", default=None, help="Ścieżka do pliku TXT z abstraktem.")
    parser.add_argument("--keywords", required=True, help="Słowa kluczowe rozdzielone średnikiem, przecinkiem lub nową linią.")
    parser.add_argument("--metadata-csv", default=None, help="Opcjonalny CSV z oficjalnymi URL-ami, profilami i metrykami.")
    parser.add_argument("--enable-web", action="store_true", help="Włącz pobieranie profili z internetu.")
    parser.add_argument(
        "--use-openalex-discovery",
        action="store_true",
        help="Spróbuj znaleźć homepage_url po ISSN przez OpenAlex, a potem pobrać profil z tej strony.",
    )
    parser.add_argument("--output-dir", default=".", help="Folder wyjściowy dla DOCX.")
    parser.add_argument("--output-file", default=None, help="Nazwa pliku wynikowego DOCX. Domyślnie tworzona z tytułu artykułu.")
    parser.add_argument("--top-n", type=int, default=DEFAULT_TOP_N, help="Liczba najlepszych czasopism w sekcji podsumowania.")
    parser.add_argument("--max-journals", type=int, default=None, help="Opcjonalny limit liczby czasopism do przetworzenia, użyteczny do testów.")
    parser.add_argument("--web-delay", type=float, default=DEFAULT_WEB_DELAY_SECONDS, help="Opóźnienie między zapytaniami HTTP w sekundach.")
    return parser


def run_stage2(
    *,
    stage1_docx_paths: list[str],
    article_title: str,
    abstract: str | None = None,
    abstract_file: str | None = None,
    keywords_raw: str,
    metadata_csv: str | None = None,
    enable_web: bool = False,
    use_openalex_discovery: bool = False,
    output_dir: str = ".",
    output_file: str | None = None,
    top_n: int = DEFAULT_TOP_N,
    max_journals: int | None = None,
    web_delay: float = DEFAULT_WEB_DELAY_SECONDS,
) -> tuple[Path, list[JournalCandidate]]:
    """
    Uruchamia Stage 2 programowo i zwraca:
    - ścieżkę do wygenerowanego DOCX,
    - posortowaną listę kandydatów (malejąco po score).
    """
    abstract_value = read_text_argument(abstract, abstract_file, "abstrakt")
    keywords = parse_keywords(keywords_raw)

    if not clean_cell(article_title):
        raise JournalFinderStage2Error("Brakuje tytułu artykułu.")
    if not abstract_value:
        raise JournalFinderStage2Error("Brakuje abstraktu artykułu.")
    if not keywords:
        raise JournalFinderStage2Error("Brakuje słów kluczowych.")

    candidates = load_stage1_docx_files([Path(path) for path in stage1_docx_paths])
    if max_journals is not None:
        if max_journals <= 0:
            raise JournalFinderStage2Error("--max-journals musi być większe od zera.")
        candidates = candidates[:max_journals]

    metadata_path = Path(metadata_csv) if metadata_csv else None
    by_issn, by_title = load_metadata_csv(metadata_path)
    for candidate in candidates:
        apply_metadata(candidate, by_issn, by_title)

    any_profile_from_csv = any(candidate.profile_text for candidate in candidates)
    any_profile_url = any(candidate.profile_url for candidate in candidates)

    if not enable_web and not any_profile_from_csv:
        raise JournalFinderStage2Error(
            "Stage 2 wymaga dostępu do internetu albo lokalnego metadata CSV z profile_text. "
            "Uruchom z --enable-web lub podaj --metadata-csv z oficjalnymi profilami czasopism."
        )

    if enable_web:
        session = build_session()
        for candidate in candidates:
            enrich_candidate_with_web(
                session=session,
                candidate=candidate,
                use_openalex_discovery=use_openalex_discovery,
                web_delay=web_delay,
            )
            time.sleep(web_delay)
    elif any_profile_url and not any_profile_from_csv:
        raise JournalFinderStage2Error(
            "W CSV są URL-e, ale pobieranie profili wymaga --enable-web."
        )

    article_text = "\n".join([article_title, abstract_value, "\n".join(keywords)])
    calculate_scores(candidates, article_text, keywords)

    target_output_dir = Path(output_dir)
    output_name = output_file or make_output_filename(article_title)
    output_path = target_output_dir / output_name
    export_results_docx(
        candidates=candidates,
        article_title=article_title,
        abstract=abstract_value,
        keywords=keywords,
        output_path=output_path,
        top_n=top_n,
    )

    sorted_candidates = sorted(candidates, key=lambda item: item.score, reverse=True)
    return output_path, sorted_candidates


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()

    try:
        output_path, sorted_candidates = run_stage2(
            stage1_docx_paths=args.stage1_docx,
            article_title=args.article_title,
            abstract=args.abstract,
            abstract_file=args.abstract_file,
            keywords_raw=args.keywords,
            metadata_csv=args.metadata_csv,
            enable_web=args.enable_web,
            use_openalex_discovery=args.use_openalex_discovery,
            output_dir=args.output_dir,
            output_file=args.output_file,
            top_n=args.top_n,
            max_journals=args.max_journals,
            web_delay=args.web_delay,
        )
        print(f"Przeanalizowano czasopism: {len(sorted_candidates)}")
        print(f"Wygenerowano DOCX: {output_path}")
        print("Najlepsze dopasowania:")
        for idx, candidate in enumerate(sorted_candidates[: args.top_n], start=1):
            print(f"{idx}. {candidate.title} — {candidate.score}/10")
        return 0

    except JournalFinderStage2Error as exc:
        print(f"Błąd: {exc}", file=sys.stderr)
        return 1
    except KeyboardInterrupt:
        print("\nPrzerwano działanie programu.", file=sys.stderr)
        return 130


if __name__ == "__main__":
    raise SystemExit(main())

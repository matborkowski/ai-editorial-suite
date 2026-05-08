from __future__ import annotations

import argparse
import math
import re
import sys
import unicodedata
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt


DEFAULT_SHEET_NAME = "Czasopisma _nauk"
DEFAULT_OUTPUT_BASENAME = "JournalFinder_stage1_results"
DEFAULT_MAX_ROWS_PER_DOCX = 700

# Struktura pliku ministerialnego 2024:
# wiersz 1: nazwy dyscyplin od kolumny J,
# wiersz 2: nazwy kolumn podstawowych i kody dyscyplin,
# wiersz 3+: dane czasopism.
ROW_DISCIPLINE_NAMES = 0
ROW_HEADERS_AND_CODES = 1
ROW_DATA_START = 2

COL_TITLE_1 = 2
COL_ISSN_1 = 3
COL_EISSN_1 = 4
COL_TITLE_2 = 5
COL_ISSN_2 = 6
COL_EISSN_2 = 7
COL_POINTS = 8
COL_FIRST_DISCIPLINE = 9


class JournalFinderError(Exception):
    """Kontrolowany błąd biznesowy Journal Finder."""


def normalize_text(value: object) -> str:
    """Normalizacja tekstu do porównań: małe litery, bez polskich znaków, bez nadmiarowych spacji."""
    if value is None:
        return ""
    text = str(value).strip().casefold()
    text = unicodedata.normalize("NFKD", text)
    text = "".join(char for char in text if not unicodedata.combining(char))
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def is_marked(value: object) -> bool:
    """W pliku dyscypliny oznaczone są znakiem 'x'."""
    return str(value).strip().casefold() == "x"


def clean_cell(value: object) -> str:
    """Zwraca bezpieczny tekst z komórki Excela."""
    if value is None:
        return ""
    if isinstance(value, float) and math.isnan(value):
        return ""
    text = str(value).strip()
    if text.lower() == "nan":
        return ""
    return text


def join_unique(values: Iterable[object]) -> str:
    """Łączy niepuste wartości, usuwając duplikaty przy zachowaniu kolejności."""
    result: list[str] = []
    seen: set[str] = set()
    for value in values:
        text = clean_cell(value)
        key = normalize_text(text)
        if text and key not in seen:
            seen.add(key)
            result.append(text)
    return "; ".join(result)


def load_ministry_file(xlsx_path: Path, sheet_name: str = DEFAULT_SHEET_NAME) -> tuple[pd.DataFrame, dict[str, int]]:
    """
    Wczytuje plik ministerialny i zwraca:
    - dataframe z danymi czasopism,
    - mapę: nazwa dyscypliny -> indeks kolumny w dataframe.
    """
    if not xlsx_path.exists():
        raise JournalFinderError(f"Nie znaleziono pliku: {xlsx_path}")

    try:
        raw = pd.read_excel(xlsx_path, sheet_name=sheet_name, header=None, dtype=object)
    except ValueError as exc:
        raise JournalFinderError(
            f"Nie znaleziono arkusza '{sheet_name}' w pliku {xlsx_path}."
        ) from exc

    if raw.shape[0] <= ROW_DATA_START or raw.shape[1] <= COL_FIRST_DISCIPLINE:
        raise JournalFinderError("Plik ma nieoczekiwaną strukturę — za mało wierszy lub kolumn.")

    discipline_cols: dict[str, int] = {}
    for col_idx in range(COL_FIRST_DISCIPLINE, raw.shape[1]):
        discipline_name = clean_cell(raw.iat[ROW_DISCIPLINE_NAMES, col_idx])
        if discipline_name:
            discipline_cols[discipline_name] = col_idx

    if not discipline_cols:
        raise JournalFinderError("Nie znaleziono kolumn z dyscyplinami naukowymi.")

    data = raw.iloc[ROW_DATA_START:].copy()
    data = data.reset_index(drop=True)
    return data, discipline_cols


def resolve_disciplines(raw_input: str, available_disciplines: list[str]) -> list[str]:
    """
    Zamienia tekst użytkownika na oficjalne nazwy dyscyplin z pliku.

    Obsługuje:
    - dokładne nazwy,
    - nazwy wpisane bez polskich znaków,
    - kilka dyscyplin rozdzielonych średnikiem lub nową linią,
    - częściowe nazwy, jeśli pasują jednoznacznie.
    """
    raw_input = raw_input.strip()
    if not raw_input:
        raise JournalFinderError("Nie podano żadnej dyscypliny.")

    norm_to_original = {normalize_text(name): name for name in available_disciplines}
    normalized_whole = normalize_text(raw_input)

    # Jedna dyscyplina podana dokładnie, także jeśli zawiera przecinki.
    if normalized_whole in norm_to_original:
        return [norm_to_original[normalized_whole]]

    # Jeżeli użytkownik wkleił pełne nazwy kilku dyscyplin w jednym tekście,
    # spróbuj wykryć je jako podciągi. To chroni dyscypliny zawierające przecinki.
    detected: list[str] = []
    remaining = f" {normalized_whole} "
    for norm_name, original_name in sorted(norm_to_original.items(), key=lambda item: len(item[0]), reverse=True):
        pattern = f" {norm_name} "
        if pattern in remaining:
            detected.append(original_name)
            remaining = remaining.replace(pattern, " ")

    leftovers = re.sub(r"[;,]+", " ", remaining)
    leftovers = re.sub(r"\s+", " ", leftovers).strip()
    if detected and not leftovers:
        return sorted(detected, key=lambda name: available_disciplines.index(name))

    # Standardowe rozdzielanie. Preferowany separator: średnik lub nowa linia.
    if ";" in raw_input or "\n" in raw_input:
        parts = [part.strip() for part in re.split(r"[;\n]+", raw_input) if part.strip()]
    else:
        # Zostawiamy obsługę przecinka, bo tak działał prototyp CustomGPT.
        # Dla dyscyplin zawierających przecinek lepiej używać średnika albo trybu interaktywnego.
        parts = [part.strip() for part in raw_input.split(",") if part.strip()]

    resolved: list[str] = []
    for part in parts:
        norm_part = normalize_text(part)
        if norm_part in norm_to_original:
            candidate = norm_to_original[norm_part]
        else:
            partial_matches = [
                original
                for norm, original in norm_to_original.items()
                if norm_part in norm or norm in norm_part
            ]
            if len(partial_matches) == 1:
                candidate = partial_matches[0]
            elif not partial_matches:
                suggestions = suggest_disciplines(part, available_disciplines)
                raise JournalFinderError(
                    f"Nie rozpoznano dyscypliny: '{part}'."
                    + (f" Podobne nazwy: {', '.join(suggestions)}." if suggestions else "")
                )
            else:
                raise JournalFinderError(
                    f"Dyscyplina '{part}' jest niejednoznaczna. Możliwe dopasowania: "
                    + "; ".join(partial_matches)
                )

        if candidate not in resolved:
            resolved.append(candidate)

    if not resolved:
        raise JournalFinderError("Nie udało się rozpoznać żadnej dyscypliny.")
    return resolved


def suggest_disciplines(query: str, available_disciplines: list[str], limit: int = 5) -> list[str]:
    """Proste sugestie dla błędnie wpisanej dyscypliny."""
    import difflib

    normalized = normalize_text(query)
    candidates = {normalize_text(name): name for name in available_disciplines}
    close_norms = difflib.get_close_matches(normalized, list(candidates.keys()), n=limit, cutoff=0.45)
    return [candidates[norm] for norm in close_norms]


def get_row_disciplines(row: pd.Series, discipline_cols: dict[str, int]) -> list[str]:
    """Zwraca listę dyscyplin przypisanych do danego czasopisma."""
    return [name for name, col_idx in discipline_cols.items() if is_marked(row.iloc[col_idx])]


def filter_journals(
    data: pd.DataFrame,
    discipline_cols: dict[str, int],
    requested_disciplines: list[str],
    min_points: int,
    max_points: int,
) -> pd.DataFrame:
    """Filtruje czasopisma zawierające wszystkie wskazane dyscypliny i mieszczące się w zakresie punktów."""
    if min_points > max_points:
        raise JournalFinderError("Minimalna liczba punktów nie może być większa od maksymalnej.")

    requested_norms = {normalize_text(name) for name in requested_disciplines}
    rows: list[dict[str, object]] = []

    for _, row in data.iterrows():
        title = clean_cell(row.iloc[COL_TITLE_1]) or clean_cell(row.iloc[COL_TITLE_2])
        if not title:
            continue

        points = pd.to_numeric(row.iloc[COL_POINTS], errors="coerce")
        if pd.isna(points):
            continue
        points_int = int(points)
        if not (min_points <= points_int <= max_points):
            continue

        row_disciplines = get_row_disciplines(row, discipline_cols)
        row_disciplines_norms = {normalize_text(name) for name in row_disciplines}
        if not requested_norms.issubset(row_disciplines_norms):
            continue

        issn = join_unique([
            row.iloc[COL_ISSN_1],
            row.iloc[COL_EISSN_1],
            row.iloc[COL_ISSN_2],
            row.iloc[COL_EISSN_2],
        ])

        rows.append(
            {
                "tytuł czasopisma": title,
                "liczba punktów": points_int,
                "dyscypliny naukowe": "; ".join(row_disciplines),
                "ISSN": issn,
            }
        )

    result = pd.DataFrame(rows, columns=["tytuł czasopisma", "liczba punktów", "dyscypliny naukowe", "ISSN"])
    if not result.empty:
        result = result.sort_values(by=["liczba punktów", "tytuł czasopisma"], ascending=[False, True])
        result = result.reset_index(drop=True)
    return result


def configure_landscape_document(document: Document) -> None:
    """Ustawia dokument DOCX w orientacji poziomej z rozsądnymi marginesami."""
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9)


def add_results_table(document: Document, result: pd.DataFrame) -> None:
    """Dodaje tabelę wynikową do dokumentu DOCX."""
    columns = list(result.columns)
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for idx, column_name in enumerate(columns):
        paragraph = header_cells[idx].paragraphs[0]
        run = paragraph.add_run(column_name)
        run.bold = True
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for _, record in result.iterrows():
        row_cells = table.add_row().cells
        for idx, column_name in enumerate(columns):
            value = clean_cell(record[column_name])
            row_cells[idx].text = value
            row_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # Orientacyjne szerokości kolumn. Word/LibreOffice i tak mogą je dopasować.
    widths = [Cm(7.5), Cm(2.2), Cm(16.0), Cm(5.0)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def export_docx_parts(
    result: pd.DataFrame,
    requested_disciplines: list[str],
    min_points: int,
    max_points: int,
    output_dir: Path,
    basename: str = DEFAULT_OUTPUT_BASENAME,
    max_rows_per_docx: int = DEFAULT_MAX_ROWS_PER_DOCX,
) -> list[Path]:
    """Eksportuje wyniki do jednego lub wielu plików DOCX."""
    if result.empty:
        return []
    if max_rows_per_docx <= 0:
        raise JournalFinderError("max_rows_per_docx musi być większe od zera.")

    output_dir.mkdir(parents=True, exist_ok=True)
    total_rows = len(result)
    part_count = math.ceil(total_rows / max_rows_per_docx)
    output_paths: list[Path] = []

    for part_idx in range(part_count):
        start = part_idx * max_rows_per_docx
        end = min(start + max_rows_per_docx, total_rows)
        part_df = result.iloc[start:end].copy()

        document = Document()
        configure_landscape_document(document)

        document.add_heading("Journal Finder – Stage 1 results", level=1)
        summary = document.add_paragraph()
        summary.add_run("Dyscypliny: ").bold = True
        summary.add_run("; ".join(requested_disciplines))
        summary.add_run("\nZakres punktów: ").bold = True
        summary.add_run(f"{min_points}–{max_points}")
        summary.add_run("\nLiczba wyników: ").bold = True
        summary.add_run(str(total_rows))
        if part_count > 1:
            summary.add_run("\nCzęść: ").bold = True
            summary.add_run(f"{part_idx + 1}/{part_count}, rekordy {start + 1}–{end}")

        add_results_table(document, part_df)

        if part_count == 1:
            output_path = output_dir / f"{basename}.docx"
        else:
            output_path = output_dir / f"{basename}_part{part_idx + 1}.docx"

        document.save(output_path)
        output_paths.append(output_path)

    return output_paths


def prompt_interactively() -> tuple[str, int, int]:
    """Tryb interaktywny, gdy użytkownik nie poda argumentów CLI."""
    print("Podaj dyscypliny naukowe.")
    print("Najbezpieczniej: jedna dyscyplina w jednej linii. Pusta linia kończy wpisywanie.")
    lines: list[str] = []
    while True:
        line = input("Dyscyplina: ").strip()
        if not line:
            break
        lines.append(line)

    disciplines_raw = "\n".join(lines)
    min_points = int(input("Minimalna liczba punktów: ").strip())
    max_points = int(input("Maksymalna liczba punktów: ").strip())
    return disciplines_raw, min_points, max_points


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Journal Finder Stage 1 — filtrowanie czasopism z ministerialnego wykazu 2024."
    )
    parser.add_argument(
        "--file",
        default="Wykaz czasopism naukowych 2024.xlsx",
        help="Ścieżka do pliku XLSX z wykazem czasopism.",
    )
    parser.add_argument(
        "--sheet",
        default=DEFAULT_SHEET_NAME,
        help=f"Nazwa arkusza z czasopismami. Domyślnie: {DEFAULT_SHEET_NAME!r}.",
    )
    parser.add_argument(
        "--disciplines",
        default=None,
        help="Dyscypliny naukowe. Kilka dyscyplin najlepiej rozdzielić średnikiem.",
    )
    parser.add_argument("--min-points", type=int, default=None, help="Minimalna liczba punktów.")
    parser.add_argument("--max-points", type=int, default=None, help="Maksymalna liczba punktów.")
    parser.add_argument(
        "--output-dir",
        default=".",
        help="Folder wyjściowy dla plików DOCX.",
    )
    parser.add_argument(
        "--max-rows-per-docx",
        type=int,
        default=DEFAULT_MAX_ROWS_PER_DOCX,
        help="Maksymalna liczba rekordów w jednym pliku DOCX. Przy większej liczbie wyników plik zostanie podzielony.",
    )
    parser.add_argument(
        "--list-disciplines",
        action="store_true",
        help="Wypisz dostępne dyscypliny i zakończ.",
    )
    return parser


def run_stage1(
    *,
    file_path: str,
    sheet_name: str = DEFAULT_SHEET_NAME,
    disciplines_raw: str,
    min_points: int,
    max_points: int,
    output_dir: str = ".",
    max_rows_per_docx: int = DEFAULT_MAX_ROWS_PER_DOCX,
    output_basename: str = DEFAULT_OUTPUT_BASENAME,
) -> tuple[list[Path], int, list[str]]:
    """
    Uruchamia Stage 1 programowo i zwraca:
    - listę wygenerowanych plików DOCX,
    - liczbę wszystkich znalezionych czasopism,
    - listę rozpoznanych dyscyplin.
    """
    xlsx_path = Path(file_path)
    target_output_dir = Path(output_dir)

    data, discipline_cols = load_ministry_file(xlsx_path, sheet_name)
    available_disciplines = list(discipline_cols.keys())
    requested_disciplines = resolve_disciplines(disciplines_raw, available_disciplines)

    result = filter_journals(
        data=data,
        discipline_cols=discipline_cols,
        requested_disciplines=requested_disciplines,
        min_points=min_points,
        max_points=max_points,
    )
    if result.empty:
        return [], 0, requested_disciplines

    output_paths = export_docx_parts(
        result=result,
        requested_disciplines=requested_disciplines,
        min_points=min_points,
        max_points=max_points,
        output_dir=target_output_dir,
        basename=output_basename,
        max_rows_per_docx=max_rows_per_docx,
    )
    return output_paths, len(result), requested_disciplines


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()

    try:
        data, discipline_cols = load_ministry_file(Path(args.file), args.sheet)
        available_disciplines = list(discipline_cols.keys())

        if args.list_disciplines:
            print("Dostępne dyscypliny naukowe:")
            for name in available_disciplines:
                print(f"- {name}")
            return 0

        if args.disciplines is None or args.min_points is None or args.max_points is None:
            disciplines_raw, min_points, max_points = prompt_interactively()
        else:
            disciplines_raw = args.disciplines
            min_points = args.min_points
            max_points = args.max_points

        output_paths, total_count, requested_disciplines = run_stage1(
            file_path=args.file,
            sheet_name=args.sheet,
            disciplines_raw=disciplines_raw,
            min_points=min_points,
            max_points=max_points,
            output_dir=args.output_dir,
            max_rows_per_docx=args.max_rows_per_docx,
        )
        if not output_paths:
            print("Brak czasopism spełniających podane kryteria. Nie wygenerowano DOCX.")
            return 0

        print(f"Znaleziono czasopism: {total_count}")
        print("Wygenerowano pliki DOCX:")
        for path in output_paths:
            print(f"- {path}")
        print("\nCzy chcesz dopasować swój artykuł do czasopism z tej listy (Stage 2)?")
        return 0

    except JournalFinderError as exc:
        print(f"Błąd: {exc}", file=sys.stderr)
        return 1
    except KeyboardInterrupt:
        print("\nPrzerwano działanie programu.", file=sys.stderr)
        return 130


if __name__ == "__main__":
    raise SystemExit(main())

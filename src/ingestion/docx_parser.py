import re
from docx import Document
from docx.oxml.ns import qn

def extract_manuscript(path: str) -> dict:
    doc = Document(path)
    
    result = {
        "title": "",
        "abstract": "",
        "keywords": [],
        "sections": {},
        "full_text": "",
        "figures": [],
        "tables": [],
    }
    
    current_section = None
    section_text = []
    full_text_parts = []
    
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]
        
        # --- TABLES ---
        if tag == "tbl":
            table_index = len(result["tables"]) + 1
            placeholder = f"[TABLE {table_index}]"
            caption = _extract_table_caption(element)
            result["tables"].append({
                "index": table_index,
                "placeholder": placeholder,
                "caption": caption,
            })
            if current_section:
                section_text.append(placeholder)
                if caption:
                    section_text.append(caption)
            full_text_parts.append(placeholder)
            continue
        
        if tag != "p":
            continue
            
        # --- PARAGRAPHS ---
        from docx.oxml import OxmlElement
        para_xml = element
        
        # CHECK FOR IMAGES
        if _contains_image(element):
            fig_index = len(result["figures"]) + 1
            placeholder = f"[FIGURE {fig_index}]"
            result["figures"].append({
                "index": fig_index,
                "placeholder": placeholder,
            })
            if current_section:
                section_text.append(placeholder)
            full_text_parts.append(placeholder)
            continue
        
        # TAKE SECTION TEXT
        text = "".join(
            node.text for node in element.iter()
            if node.tag.split("}")[-1] == "t" and node.text
        ).strip()
        
        if not text:
            continue
        
        # SECTION STYLE
        style = ""
        pPr = element.find(f".//{{{element.nsmap.get('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')}}}pStyle")
        if pPr is not None:
            style = pPr.get(f"{{{element.nsmap.get('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')}}}val", "").lower()
        
        # SEARCH FOR FIGURES/TABLES DESCRIPTION
        if _is_caption(text):
            if current_section:
                section_text.append(text)
            full_text_parts.append(text)
            continue
        
        # CHECK SECTION HEADING
        if "heading" in style or _is_section_heading(text):
            if current_section and section_text:
                result["sections"][current_section] = "\n".join(section_text)
            current_section = text.strip()
            section_text = []
            full_text_parts.append(text)
            continue
        
        # TITLE
        if not result["title"] and _looks_like_title(text, style):
            result["title"] = text
            full_text_parts.append(text)
            continue
        
        # ABSTRACT
        if not result["abstract"] and current_section and "abstract" in current_section.lower():
            result["abstract"] += text + " "
        
        # KEYWORDS
        if current_section and "keyword" in current_section.lower():
            result["keywords"] = [k.strip() for k in re.split(r"[;,]", text) if k.strip()]
        
        # ADD TO CURRENT SECTION
        if current_section:
            section_text.append(text)
        full_text_parts.append(text)
    
    # SAVE LAST SECTION
    if current_section and section_text:
        result["sections"][current_section] = "\n".join(section_text)
    
    result["full_text"] = "\n".join(full_text_parts)
    result["abstract"] = result["abstract"].strip()
    
    return result


def extract_paragraphs(path: str) -> list[str]:
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _contains_image(element) -> bool:
    ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    drawing_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    return (
        element.find(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline") is not None
        or element.find(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor") is not None
    )


def _extract_table_caption(element) -> str:
    """Try to find a caption paragraph immediately after a table."""
    return ""  # rozszerzyć jeśli potrzebne


def _is_caption(text: str) -> bool:
    """Detect figure and table captions."""
    patterns = [
        r"^(Fig\.?|Figure)\s*\d+",
        r"^Table\s*\d+",
        r"^\[Figure\s*\d+",
        r"^\[Table\s*\d+",
        r"^<Figure\s*\d+",
        r"^<Table\s*\d+",
    ]
    return any(re.match(p, text, re.IGNORECASE) for p in patterns)


def _is_section_heading(text: str) -> bool:
    """Detect common academic section headings."""
    headings = [
        "abstract", "introduction", "background", "literature",
        "methods", "materials", "methodology", "experimental",
        "results", "discussion", "conclusion", "conclusions",
        "acknowledgments", "acknowledgements", "references",
        "appendix", "keywords", "funding",
    ]
    return text.strip().lower().rstrip(".") in headings


def _looks_like_title(text: str, style: str) -> bool:
    return "title" in style or (len(text) > 20 and len(text) < 300 and text[0].isupper())